#!/usr/bin/env python3
"""
generate_2568_docx.py
ดึงข้อมูลปีการศึกษา 2568 และสร้างไฟล์เอกสารรายงาน (.docx) 
จัดระเบียบตามชั้นปีการศึกษาในห้อง data/exports/AY2568/
ใช้รูปแบบเอกสารข้าราชการกองทัพอากาศ (RTAF) ปรับแต่งขอบกระดาษและฟอนต์ขนาด 16pt
และแทรก Zero-Width Space (U+200B) เพื่อการตัดบรรทัดที่สวยงาม
"""

import json
import os
import re
import shutil
from pathlib import Path
from datetime import datetime

# Safe imports for docx and pythainlp
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from pythainlp import word_tokenize
    DOCX_SUPPORTED = True
except ImportError as e:
    print(f"❌ Dependencies missing: {e}")
    print("Please make sure python-docx and pythainlp are installed.")
    DOCX_SUPPORTED = False

BASE_DIR = Path(__file__).parent.parent
DATA_DIR = BASE_DIR / "data"
EXPORTS_DIR = DATA_DIR / "exports" / "AY2568"

STUDENTS_JSON = BASE_DIR / "frontend" / "data" / "students.json"
DEEDS_JSON = BASE_DIR / "frontend" / "data" / "deeds.json"

ZWS = "\u200b"  # Zero-Width Space

CATEGORIES_NAME_MAP = {
    1: 'บริจาคโลหิต/เกล็ดเลือด/พลาสมา',
    2: 'โครงการภายนอก (คำสั่ง วพอ.)',
    3: 'ช่วยเหลืองานภายใน วพอ.',
    4: 'เข้าอบรมที่ วพอ. จัดให้',
    5: 'ช่วยงานหน่วยงาน/ชุมชน/มูลนิธิ',
    6: 'ทำนุบำรุงศาสนสถาน',
    7: 'งานฟรีทั่วไป',
    8: 'กิจกรรมจงรักภักดีต่อสถาบัน',
    9: 'ชม. ที่สมควรได้รับ (บทบาทพิเศษ)'
}

def insert_zwsp(text, engine="newmm"):
    """แทรก Zero-Width Space ระหว่างคำไทยเพื่อการตัดคำใน MS Word"""
    if not text:
        return ""
    thai_pattern = re.compile(r'([\u0E00-\u0E7F]+)')
    parts = thai_pattern.split(str(text))
    result = []
    for part in parts:
        if thai_pattern.fullmatch(part):
            tokens = word_tokenize(part, engine=engine)
            result.append(ZWS.join(tokens))
        else:
            result.append(part)
    return "".join(result)

def get_category_name(category_id):
    return CATEGORIES_NAME_MAP.get(int(category_id), 'กิจกรรมจิตอาสา')

def generate_report(student, deeds, output_path):
    """สร้างไฟล์ docx สำหรับรายงานส่วนบุคคลของนักเรียนแต่ละคน"""
    doc = Document()
    
    # 1. Page Setup (RTAF Standard: Left 3.0cm, Right 2.0cm, Top 2.5cm, Bottom 2.0cm)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    
    # 2. Font Setup (RTAF Standard: TH Sarabun New size 16)
    font_name = "TH Sarabun New"
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(16)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # Report Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(insert_zwsp("รายงานสรุปการบำเพ็ญประโยชน์จิตอาสา"))
    run_title.font.name = font_name
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run(insert_zwsp("วิทยาลัยพยาบาลทหารอากาศ ปีการศึกษา 2568"))
    run_sub.font.name = font_name
    run_sub.font.size = Pt(16)
    run_sub.font.bold = True
    
    # Section 1: Personal Info
    p_profile_title = doc.add_paragraph()
    p_profile_title.paragraph_format.space_after = Pt(6)
    run_profile_title = p_profile_title.add_run(insert_zwsp("๑. ข้อมูลประจำตัวนักเรียนพยาบาลทหารอากาศ"))
    run_profile_title.font.name = font_name
    run_profile_title.font.size = Pt(16)
    run_profile_title.font.bold = True
    
    fullname = f"{student.get('rank', 'นพอ.')} {student.get('first_name', '')} {student.get('last_name', '')}"
    total_hours = sum(float(d.get('hours', 0)) for d in deeds)
    passed = total_hours >= 50
    
    p_info = doc.add_paragraph()
    p_info.paragraph_format.left_indent = Cm(1.0)
    p_info.paragraph_format.space_after = Pt(12)
    
    info_text = (
        f"ชื่อ-สกุล: {fullname}\n"
        f"เลขประจำตัวนักเรียน: {student.get('student_id', '')}   ชั้นปี: {student.get('year_level', '')} (รุ่น {student.get('class_year', '')})\n"
        f"ชั่วโมงจิตอาสาสะสม ปีการศึกษา 2568: {total_hours:.1f} ชั่วโมง (จากเกณฑ์ขั้นต่ำ 50 ชั่วโมง)\n"
        f"ผลการประเมินชั่วโมงจิตอาสา: {'ผ่านเกณฑ์การสะสมชั่วโมง' if passed else 'ยังไม่ผ่านเกณฑ์การสะสมชั่วโมง'}"
    )
    p_info.add_run(insert_zwsp(info_text))
    
    # Section 2: Deeds Table
    p_deeds_title = doc.add_paragraph()
    p_deeds_title.paragraph_format.space_before = Pt(12)
    p_deeds_title.paragraph_format.space_after = Pt(6)
    run_deeds_title = p_deeds_title.add_run(insert_zwsp("๒. ประวัติการบำเพ็ญประโยชน์จิตอาสาที่ได้รับการอนุมัติ"))
    run_deeds_title.font.name = font_name
    run_deeds_title.font.size = Pt(16)
    run_deeds_title.font.bold = True
    
    if not deeds:
        p_empty = doc.add_paragraph()
        p_empty.paragraph_format.left_indent = Cm(1.0)
        p_empty.add_run(insert_zwsp("- ไม่พบประวัติกิจกรรมจิตอาสาที่ได้รับการอนุมัติในระบบ -"))
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = insert_zwsp("ลำดับ")
        hdr_cells[1].text = insert_zwsp("วันที่ทำกิจกรรม")
        hdr_cells[2].text = insert_zwsp("ประเภทจิตอาสา")
        hdr_cells[3].text = insert_zwsp("รายละเอียดกิจกรรม")
        hdr_cells[4].text = insert_zwsp("ชั่วโมง")
        
        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.name = font_name
                run.font.bold = True
                run.font.size = Pt(14)
                
        # Rows
        for idx, d in enumerate(deeds, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = d.get('activityDate', '')
            
            category_id = d.get('categoryId', 0)
            category_name = get_category_name(category_id)
            row_cells[2].text = insert_zwsp(category_name)
            row_cells[3].text = insert_zwsp(d.get('description', ''))
            row_cells[4].text = f"{float(d.get('hours', 0)):.1f}"
            
            for col_idx, cell in enumerate(row_cells):
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                if col_idx in (0, 1, 4):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = font_name
                    run.font.size = Pt(13)
                    
        # Widths: Total 16.0cm matching margins (21 - 3 - 2 = 16)
        widths = [Cm(1.2), Cm(2.8), Cm(3.8), Cm(6.5), Cm(1.7)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                
    # Signatures
    p_sig_space = doc.add_paragraph()
    p_sig_space.paragraph_format.space_before = Pt(40)
    
    table_sig = doc.add_table(rows=1, cols=2)
    table_sig.autofit = False
    
    sig_cells = table_sig.rows[0].cells
    sig_cells[0].width = Cm(8.0)
    sig_cells[1].width = Cm(8.0)
    
    p_sig_left = sig_cells[0].paragraphs[0]
    p_sig_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig_left.add_run(insert_zwsp(
        "ลงชื่อ...........................................................\n"
        f"( {fullname} )\n"
        "นักเรียนพยาบาลทหารอากาศ ผู้รายงาน"
    ))
    for run in p_sig_left.runs:
        run.font.name = font_name
        run.font.size = Pt(14)
        
    p_sig_right = sig_cells[1].paragraphs[0]
    p_sig_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig_right.add_run(insert_zwsp(
        "ลงชื่อ...........................................................\n"
        "( ........................................................... )\n"
        "อาจารย์ผู้รับรอง / ผู้ตรวจสอบ"
    ))
    for run in p_sig_right.runs:
        run.font.name = font_name
        run.font.size = Pt(14)
        
    # Save
    doc.save(output_path)

def main():
    if not DOCX_SUPPORTED:
        return

    print("🚀 Starting 2568 Academic Year DOCX Report Generation...")
    
    # 1. Load data
    if not STUDENTS_JSON.exists():
        print(f"❌ File not found: {STUDENTS_JSON}")
        return
    if not DEEDS_JSON.exists():
        print(f"❌ File not found: {DEEDS_JSON}")
        return
        
    with open(STUDENTS_JSON, "r", encoding="utf-8") as f:
        students = json.load(f)
    with open(DEEDS_JSON, "r", encoding="utf-8") as f:
        all_deeds = json.load(f)
        
    student_map = {s["student_id"]: s for s in students}
    
    # Recreate exports folder
    if EXPORTS_DIR.exists():
        shutil.rmtree(EXPORTS_DIR)
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    
    count = 0
    # 2. Iterate through deeds to find 2568 entries
    for student_id, deeds_list in all_deeds.items():
        # Filter deeds belonging to 2568
        deeds_2568 = [d for d in deeds_list if d.get("academicYear") == 2568 and d.get("status") == "approved"]
        if not deeds_2568:
            continue
            
        student = student_map.get(student_id)
        if not student:
            print(f"⚠️ Student ID {student_id} has 2568 deeds but no profile in students.json. Skipping...")
            continue
            
        class_year = student.get("class_year", 66)
        class_dir = EXPORTS_DIR / f"Class_{class_year}"
        class_dir.mkdir(exist_ok=True)
        
        first_name = student.get("first_name", "Unknown").replace(" ", "_")
        last_name = student.get("last_name", "Unknown").replace(" ", "_")
        filename = f"{student_id}_{first_name}_{last_name}.docx"
        output_path = class_dir / filename
        
        try:
            generate_report(student, deeds_2568, output_path)
            count += 1
            if count % 20 == 0:
                print(f"   Generated {count} reports...")
        except Exception as e:
            print(f"❌ Failed to generate report for {student_id}: {e}")
            
    print(f"\n🎉 Successfully generated {count} RTAF official reports under data/exports/AY2568/")
    
    # Create a zip of the exports for the user to easily download
    zip_output = DATA_DIR / "exports" / "AY2568_Reports"
    shutil.make_archive(str(zip_output), 'zip', str(EXPORTS_DIR))
    print(f"📦 Created zip file: {zip_output}.zip")

if __name__ == "__main__":
    main()
