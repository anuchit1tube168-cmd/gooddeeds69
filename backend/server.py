import json
import os
import time
import sys
import re
import io
import threading
import queue
from http.server import SimpleHTTPRequestHandler, HTTPServer, ThreadingHTTPServer
from urllib.parse import urlparse, parse_qs
from http.cookies import SimpleCookie

# Real-time event streams for connected clients
clients = []
clients_lock = threading.Lock()

def broadcast_event(event_type, data):
    """Send an event to all connected SSE clients."""
    with clients_lock:
        message = f"event: {event_type}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"
        for q in clients:
            try:
                q.put_nowait(message)
            except Exception:
                pass


# Try to import DOCX and pythainlp libraries safely
try:
    from pythainlp import word_tokenize
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_SUPPORTED = True
except ImportError as e:
    print(f"⚠️  DOCX generation libraries not fully available: {e}")
    DOCX_SUPPORTED = False

# Set the base directory to the project root
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FRONTEND_DIR = os.path.join(BASE_DIR, 'frontend')
RECORDS_DIR = os.path.join(BASE_DIR, 'records')
GDRIVE_DEST = os.path.expanduser('~/Library/CloudStorage/GoogleDrive-anuchit1tube168@gmail.com/ไดรฟ์ของฉัน/ระบบบันทึกความดี_วพอ_2569')

def sync_to_google_drive_bg():
    """Sync data folder to Google Drive in background thread."""
    def run_sync():
        try:
            if os.path.exists(os.path.dirname(GDRIVE_DEST)):
                os.makedirs(GDRIVE_DEST, exist_ok=True)
                import subprocess
                subprocess.run([
                    'rsync', '-av',
                    '--exclude=.git',
                    '--exclude=node_modules',
                    '--exclude=photos_backup',
                    BASE_DIR + '/',
                    GDRIVE_DEST + '/'
                ], check=False, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                print("☁️ Synced latest data to Google Drive successfully!")
        except Exception as e:
            print(f"⚠️ Google Drive sync warning: {e}")
    threading.Thread(target=run_sync, daemon=True).start()

CATEGORIES_NAME_MAP = {
    1: 'บริจาคโลหิต/เกล็ดเลือด/พลาสมา',
    2: 'โครงการภายนอก (คำสั่ง วพอ.)',
    3: 'ช่วยเหลืองานภายใน วพอ.',
    4: 'เข้าอบรมที่ วพอ. จัดให้',
    5: 'ช่วยงานหน่วยงาน/ชุมชน/มูลนิธิ',
    6: 'ทำนุบำรุงศาสนสถาน',
    7: 'งานฟรีทั่วไป',
    8: 'กิจกรรมจงรักภักดีต่อสถาบัน',
    9: 'ชม. ที่สมควรได้รับ (บทบาทพิเศษ)'
}

ZWS = "\u200b"  # Zero-Width Space

def insert_zwsp(text, engine="newmm"):
    """Insert Zero-Width Space between Thai words to allow proper Word wrapping."""
    if not text:
        return ""
    thai_pattern = re.compile(r'([\u0E00-\u0E7F]+)')
    parts = thai_pattern.split(str(text))
    result = []
    for part in parts:
        if thai_pattern.fullmatch(part):
            try:
                tokens = word_tokenize(part, engine=engine)
                result.append(ZWS.join(tokens))
            except:
                result.append(part)
        else:
            result.append(part)
    return "".join(result)

def get_category_name(category_id):
    try:
        return CATEGORIES_NAME_MAP.get(int(category_id), "อื่นๆ")
    except:
        return "อื่นๆ"

def parse_cookie_header(cookie_header):
    if not cookie_header:
        return {}
    cookie = SimpleCookie()
    try:
        cookie.load(cookie_header)
    except Exception:
        return {}
    return {key: morsel.value for key, morsel in cookie.items()}

def is_staff_role(role):
    return role in ('teacher', 'admin')

def get_deeds_for_student(student_id):
    """Retrieve all deeds for a student, merging seeded seeds with dynamic records."""
    # 1. Load imported deeds from frontend/data/deeds.json
    imported_deeds = {}
    deeds_json_path = os.path.join(BASE_DIR, 'frontend', 'data', 'deeds.json')
    if os.path.exists(deeds_json_path):
        try:
            with open(deeds_json_path, 'r', encoding='utf-8') as f:
                imported_deeds = json.load(f)
        except Exception as e:
            print(f"Error reading deeds.json: {e}")
            
    student_deeds = imported_deeds.get(student_id, [])
    
    # 2. Load dynamic deeds from records/ recursively
    dynamic_deeds = []
    if os.path.exists(RECORDS_DIR):
        for root, dirs, files in os.walk(RECORDS_DIR):
            if f"Student_{student_id}" in root:
                for file in files:
                    if file.endswith('.json'):
                        file_path = os.path.join(root, file)
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                deed = json.load(f)
                                dynamic_deeds.append(deed)
                        except Exception as e:
                            print(f"Error reading dynamic deed {file_path}: {e}")
                            
    # 3. Merge them to avoid duplicates by deed ID
    merged_map = {d['id']: d for d in student_deeds}
    for d in dynamic_deeds:
        merged_map[d['id']] = d
        
    return list(merged_map.values())

def get_all_deeds():
    """Retrieve all deeds for all students (seeded + dynamic)."""
    # 1. Load all imported deeds
    imported_deeds = {}
    deeds_json_path = os.path.join(BASE_DIR, 'frontend', 'data', 'deeds.json')
    if os.path.exists(deeds_json_path):
        try:
            with open(deeds_json_path, 'r', encoding='utf-8') as f:
                imported_deeds = json.load(f)
        except Exception as e:
            print(f"Error reading deeds.json: {e}")
            
    # 2. Walk records/ directory to find all dynamic deeds
    dynamic_deeds_by_student = {}
    if os.path.exists(RECORDS_DIR):
        for root, dirs, files in os.walk(RECORDS_DIR):
            match = re.search(r'Student_(\d+)', root)
            if match:
                student_id = match.group(1)
                if student_id not in dynamic_deeds_by_student:
                    dynamic_deeds_by_student[student_id] = []
                for file in files:
                    if file.endswith('.json'):
                        file_path = os.path.join(root, file)
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                deed = json.load(f)
                                dynamic_deeds_by_student[student_id].append(deed)
                        except Exception as e:
                            print(f"Error reading dynamic deed {file_path}: {e}")
                            
    # 3. Merge
    all_students = set(list(imported_deeds.keys()) + list(dynamic_deeds_by_student.keys()))
    result = {}
    for sid in all_students:
        student_deeds = imported_deeds.get(sid, [])
        dyn_deeds = dynamic_deeds_by_student.get(sid, [])
        merged_map = {d['id']: d for d in student_deeds}
        for d in dyn_deeds:
            merged_map[d['id']] = d
        result[sid] = list(merged_map.values())
        
    return result

def get_academic_term(date_str):
    """Return (academic_year, semester) for a given date string (YYYY-MM-DD)."""
    try:
        parts = str(date_str).split('-')
        year = int(parts[0])
        month = int(parts[1]) if len(parts) > 1 else 7
        be_year = year + 543 if year < 2400 else year
        academic_year = be_year if month >= 7 else be_year - 1
        semester = 1 if 7 <= month <= 11 else 2
        return academic_year, semester
    except Exception:
        return 2569, 1

def validate_deed_submission(deed_data, student_id):
    """Validate deed limits and prevent duplicates according to RTAFNC rules."""
    try:
        cat_id = int(deed_data.get('categoryId') or deed_data.get('category_id') or 0)
        hours = float(deed_data.get('hours', 0))
        act_date = deed_data.get('activityDate') or deed_data.get('event_date') or ''
        desc = (deed_data.get('description') or deed_data.get('title') or '').strip()
        deed_id = str(deed_data.get('id', '')).strip()

        if cat_id < 1 or cat_id > 9:
            return False, "กรุณาเลือกหมวดหมู่ความดีที่ถูกต้อง (หมวด 1 - 9)"
        if hours <= 0:
            return False, "จำนวนชั่วโมงต้องมากกว่า 0"
        if not act_date:
            return False, "กรุณาระบุวันที่ทำกิจกรรม"
        if not desc:
            return False, "กรุณากรอกรายละเอียดกิจกรรม"

        academic_year, semester = get_academic_term(act_date)

        # Per session limits
        max_session_hours = {
            1: 8.0,
            2: 8.0,
            3: 8.0,
            4: 4.0,
            5: 4.0,
            6: 1.0,
            7: 1.0,
            8: 8.0,
            9: 20.0,
        }
        if cat_id in max_session_hours and hours > max_session_hours[cat_id]:
            return False, f"⚠️ หมวดที่ {cat_id} กำหนดให้บันทึกได้ไม่เกิน {max_session_hours[cat_id]:.0f} ชม. ต่อครั้ง (คุณกรอก {hours} ชม.)"

        # Load existing deeds for student
        existing_deeds = get_deeds_for_student(student_id)
        valid_deeds = [d for d in existing_deeds if d.get('status') != 'rejected' and str(d.get('id', '')) != deed_id]

        # Duplicate check: same category and same activityDate
        for d in valid_deeds:
            d_cat = int(d.get('categoryId') or d.get('category_id') or 0)
            d_date = d.get('activityDate') or d.get('event_date') or ''
            d_desc = (d.get('description') or d.get('title') or '').strip().lower()
            if d_cat == cat_id and d_date == act_date:
                if d_desc == desc.lower() or abs(float(d.get('hours', 0)) - hours) < 0.01:
                    return False, f"⚠️ พบรายการความดีในหมวดนี้ในวันที่ {act_date} บันทึกอยู่แล้วในระบบ เพื่อป้องกันการบันทึกซ้ำ"

        # Filter deeds for the same academic year
        year_deeds = []
        sem_deeds = []
        for d in valid_deeds:
            d_cat = int(d.get('categoryId') or d.get('category_id') or 0)
            if d_cat == cat_id:
                d_date = d.get('activityDate') or d.get('event_date') or ''
                d_ay, d_sem = get_academic_term(d_date)
                if d_ay == academic_year:
                    year_deeds.append(d)
                    if d_sem == semester:
                        sem_deeds.append(d)

        year_hours = sum(float(d.get('hours', 0)) for d in year_deeds)
        sem_hours = sum(float(d.get('hours', 0)) for d in sem_deeds)

        # 1. Category 1: Blood Donation (8 hrs/session, max 4 times/yr, 90 days spacing)
        if cat_id == 1:
            if len(year_deeds) >= 4:
                return False, f"⚠️ การบริจาคโลหิตบันทึกได้ไม่เกิน 4 ครั้งต่อปีการศึกษา (ปีการศึกษา {academic_year} บันทึกครบ 4 ครั้งแล้ว)"
            
            # Spacing check: 90 days from any previous blood donation
            try:
                from datetime import datetime
                new_dt = datetime.strptime(act_date[:10], '%Y-%m-%d')
                for d in valid_deeds:
                    d_cat = int(d.get('categoryId') or d.get('category_id') or 0)
                    if d_cat == 1:
                        d_date = d.get('activityDate') or d.get('event_date') or ''
                        if len(d_date) >= 10:
                            prev_dt = datetime.strptime(d_date[:10], '%Y-%m-%d')
                            diff_days = abs((new_dt - prev_dt).days)
                            if diff_days < 90:
                                return False, f"⚠️ การบริจาคโลหิตต้องเว้นระยะห่างอย่างน้อย 3 เดือน (90 วัน) กิจกรรมนี้ห่างจากการบริจาคเมื่อ {d_date} เพียง {diff_days} วัน"
            except Exception:
                pass

        # 2. Category 6: ศาสนสถาน (max 1 hr/session, max 4 hrs/year)
        elif cat_id == 6:
            if hours > 1.0:
                return False, "⚠️ งานทำนุบำรุงศาสนสถานบันทึกได้ไม่เกิน 1 ชั่วโมงต่อครั้ง"
            if year_hours + hours > 4.0:
                rem = max(0.0, 4.0 - year_hours)
                return False, f"⚠️ งานทำนุบำรุงศาสนสถานสะสมได้ไม่เกิน 4 ชม./ปีการศึกษา (ปีนี้ใช้ไปแล้ว {year_hours:.1f} ชม., เหลือโควตา {rem:.1f} ชม.)"

        # 3. Category 7: งานฟรีทั่วไป / ช่วยผู้ปกครอง (max 1 hr/session, max 2 hrs/sem, max 4 hrs/year)
        elif cat_id == 7:
            if hours > 1.0:
                return False, "⚠️ งานฟรีทั่วไป/ช่วยงานผู้ปกครอง บันทึกได้ไม่เกิน 1 ชั่วโมงต่อครั้ง"
            if sem_hours + hours > 2.0:
                rem_sem = max(0.0, 2.0 - sem_hours)
                return False, f"⚠️ งานฟรีทั่วไป/ช่วยงานผู้ปกครอง บันทึกได้ไม่เกิน 2 ชม. ต่อภาคการศึกษา (เทอม {semester} ใช้ไปแล้ว {sem_hours:.1f} ชม., เหลือโควตา {rem_sem:.1f} ชม.)"
            if year_hours + hours > 4.0:
                rem_yr = max(0.0, 4.0 - year_hours)
                return False, f"⚠️ งานฟรีทั่วไป/ช่วยงานผู้ปกครอง สะสมได้ไม่เกิน 4 ชม./ปีการศึกษา (ปีนี้ใช้ไปแล้ว {year_hours:.1f} ชม., เหลือโควตา {rem_yr:.1f} ชม.)"

        # 4. Category 8: จงรักภักดี (max 8 hrs/year)
        elif cat_id == 8:
            if year_hours + hours > 8.0:
                rem_yr = max(0.0, 8.0 - year_hours)
                return False, f"⚠️ กิจกรรมจงรักภักดีต่อสถาบัน สะสมได้ไม่เกิน 8 ชม./ปีการศึกษา (ปีนี้ใช้ไปแล้ว {year_hours:.1f} ชม., เหลือโควตา {rem_yr:.1f} ชม.)"

        # 5. Category 9: บทบาทพิเศษ (max 20 hrs/sem)
        elif cat_id == 9:
            if sem_hours + hours > 20.0:
                return False, f"⚠️ บทบาทพิเศษ บันทึกได้ไม่เกิน 20 ชม. ต่อภาคการศึกษา (เทอม {semester} บันทึกไปแล้ว {sem_hours:.1f} ชม.)"

        return True, "Valid"
    except Exception as e:
        return True, str(e)

def generate_docx_in_memory(student_id, academic_year=2569):
    """Generate a professional A4 Word document report for a student."""
    if not DOCX_SUPPORTED:
        raise RuntimeError("docx and pythainlp libraries are not installed on this server.")

    # 1. Load student info from frontend/data/students.json
    students = []
    students_json_path = os.path.join(BASE_DIR, 'frontend', 'data', 'students.json')
    if os.path.exists(students_json_path):
        try:
            with open(students_json_path, 'r', encoding='utf-8') as f:
                students = json.load(f)
        except Exception as e:
            print(f"Error reading students.json: {e}")
            
    student = next((s for s in students if s['student_id'] == student_id), None)
    if not student:
        return None
        
    # 2. Get approved deeds
    deeds = get_deeds_for_student(student_id)
    approved_deeds = [d for d in deeds if d.get('status') == 'approved' and d.get('academicYear', 2569) == academic_year]
    approved_deeds.sort(key=lambda x: x.get('activityDate', ''))
    
    total_hours = sum(float(d.get('hours', 0)) for d in approved_deeds)
    passed = total_hours >= 50
    
    # 3. Create document
    doc = Document()
    
    # Page Setup (A4 with 2.54 cm margins)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    
    # Default Font Setup (TH Sarabun New)
    font_name = "TH Sarabun New"
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(15)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)
    
    # Report Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(insert_zwsp("รายงานสรุปการบำเพ็ญประโยชน์จิตอาสา"))
    run_title.font.name = font_name
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run(insert_zwsp(f"วิทยาลัยพยาบาลทหารอากาศ ปีการศึกษา {academic_year}"))
    run_sub.font.name = font_name
    run_sub.font.size = Pt(16)
    run_sub.font.bold = True
    
    # Student Profile Section
    p_profile_title = doc.add_paragraph()
    p_profile_title.paragraph_format.space_after = Pt(6)
    run_profile_title = p_profile_title.add_run(insert_zwsp("๑. ข้อมูลประจำตัวนักเรียนพยาบาลทหารอากาศ"))
    run_profile_title.font.name = font_name
    run_profile_title.font.size = Pt(16)
    run_profile_title.font.bold = True
    
    fullname = f"{student.get('rank', 'นพอ.')} {student.get('first_name', '')} {student.get('last_name', '')}"
    
    p_info = doc.add_paragraph()
    p_info.paragraph_format.left_indent = Cm(1.0)
    p_info.paragraph_format.space_after = Pt(6)
    
    info_text = (
        f"ชื่อ-สกุล: {fullname}\n"
        f"เลขประจำตัวนักเรียน: {student.get('student_id', '')}   ชั้นปี: {student.get('year_level', '')} (รุ่น {student.get('class_year', '')})\n"
        f"ชั่วโมงจิตอาสาสะสมได้รับการอนุมัติ: {total_hours:.1f} ชั่วโมง (จากเกณฑ์ขั้นต่ำ 50 ชั่วโมง)\n"
        f"ผลการประเมินชั่วโมงจิตอาสา: {'ผ่านเกณฑ์การสะสมชั่วโมง' if passed else 'ยังไม่ผ่านเกณฑ์การสะสมชั่วโมง'}"
    )
    p_info.add_run(insert_zwsp(info_text))
    
    # Deed Records Section
    p_deeds_title = doc.add_paragraph()
    p_deeds_title.paragraph_format.space_before = Pt(12)
    p_deeds_title.paragraph_format.space_after = Pt(6)
    run_deeds_title = p_deeds_title.add_run(insert_zwsp("๒. ประวัติการบำเพ็ญประโยชน์จิตอาสาที่ได้รับการอนุมัติ"))
    run_deeds_title.font.name = font_name
    run_deeds_title.font.size = Pt(16)
    run_deeds_title.font.bold = True
    
    if not approved_deeds:
        p_empty = doc.add_paragraph()
        p_empty.paragraph_format.left_indent = Cm(1.0)
        p_empty.add_run(insert_zwsp("- ไม่พบประวัติกิจกรรมจิตอาสาที่ได้รับการอนุมัติในระบบ -"))
    else:
        # Table of approved deeds
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = insert_zwsp("ลำดับ")
        hdr_cells[1].text = insert_zwsp("วันที่ทำกิจกรรม")
        hdr_cells[2].text = insert_zwsp("ประเภทจิตอาสา")
        hdr_cells[3].text = insert_zwsp("รายละเอียดกิจกรรม")
        hdr_cells[4].text = insert_zwsp("ชั่วโมง")
        
        # Center align headers and make bold
        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.name = font_name
                run.font.bold = True
                run.font.size = Pt(14)
                
        # Populate rows
        for idx, d in enumerate(approved_deeds, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = d.get('activityDate', '')
            
            category_id = d.get('categoryId', 0)
            category_name = get_category_name(category_id)
            row_cells[2].text = insert_zwsp(category_name)
            row_cells[3].text = insert_zwsp(d.get('description', ''))
            row_cells[4].text = f"{float(d.get('hours', 0)):.1f}"
            
            # Formats
            for col_idx, cell in enumerate(row_cells):
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                if col_idx in (0, 1, 4):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = font_name
                    run.font.size = Pt(13)
                    
        # Column Widths
        widths = [Cm(1.2), Cm(2.8), Cm(3.8), Cm(6.5), Cm(1.7)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                
    # Signatures
    p_sig_space = doc.add_paragraph()
    p_sig_space.paragraph_format.space_before = Pt(40)
    
    table_sig = doc.add_table(rows=1, cols=2)
    table_sig.autofit = False
    
    sig_cells = table_sig.rows[0].cells
    sig_cells[0].width = Cm(8.0)
    sig_cells[1].width = Cm(8.0)
    
    p_sig_left = sig_cells[0].paragraphs[0]
    p_sig_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig_left.add_run(insert_zwsp(
        "ลงชื่อ...........................................................\n"
        f"( {fullname} )\n"
        "นักเรียนพยาบาลทหารอากาศ ผู้รายงาน"
    ))
    for run in p_sig_left.runs:
        run.font.name = font_name
        run.font.size = Pt(14)
        
    p_sig_right = sig_cells[1].paragraphs[0]
    p_sig_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig_right.add_run(insert_zwsp(
        "ลงชื่อ...........................................................\n"
        "( ........................................................... )\n"
        "อาจารย์ผู้รับรอง / ผู้ตรวจสอบ"
    ))
    for run in p_sig_right.runs:
        run.font.name = font_name
        run.font.size = Pt(14)
        
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def load_students_map():
    for p in [os.path.join(BASE_DIR, 'data', 'students.json'), os.path.join(BASE_DIR, 'frontend', 'data', 'students.json')]:
        if os.path.exists(p):
            try:
                with open(p, 'r', encoding='utf-8') as f:
                    return {str(s['student_id']).strip(): s for s in json.load(f)}
            except Exception:
                pass
    return {}

def save_or_update_deed_in_db(student_id, deed_data):
    """Persist deed into records/ directory and sync into deeds.json and deeds_data.js."""
    student_id = str(student_id).strip()
    deed_id = str(deed_data.get('id', '')).strip()
    if not student_id or not deed_id:
        return False

    academic_year = deed_data.get('academicYear', 2569)
    student = deed_data.get('student', {})
    class_year = student.get('class_year') or (student_id[:2] if len(student_id) >= 2 else '69')
    category_id = deed_data.get('categoryId') or deed_data.get('category_id') or 7

    # Ensure deed has real student identity info
    if not deed_data.get('student_name'):
        s_map = load_students_map()
        s = s_map.get(student_id)
        if s:
            deed_data['student_name'] = f"{s.get('rank', 'นพอ.')} {s.get('first_name', '')} {s.get('last_name', '')}".strip()
            deed_data['studentName'] = deed_data['student_name']
            deed_data['student_rank'] = s.get('rank', 'นพอ.')
            deed_data['student_first_name'] = s.get('first_name', '')
            deed_data['student_last_name'] = s.get('last_name', '')
            deed_data['class_year'] = str(s.get('class_year', class_year))
            if s.get('position'): deed_data['student_position'] = s.get('position')

    # 1. Save directly into records/
    target_dir = os.path.join(
        RECORDS_DIR,
        f"AY{academic_year}",
        f"Class_{class_year}",
        f"Student_{student_id}",
        f"Category_{category_id}"
    )
    os.makedirs(target_dir, exist_ok=True)
    filepath = os.path.join(target_dir, f"{deed_id}.json")
    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(deed_data, f, ensure_ascii=False, indent=4)
    except Exception as e:
        print(f"⚠️ Error saving record {filepath}: {e}")

    # 2. Update data/deeds.json and frontend/data/deeds.json
    deeds_file = os.path.join(BASE_DIR, 'data', 'deeds.json')
    data = {}
    if os.path.exists(deeds_file):
        try:
            with open(deeds_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception:
            data = {}

    if student_id not in data:
        data[student_id] = []

    # Find existing or append
    found = False
    for i, d in enumerate(data[student_id]):
        if str(d.get('id')) == deed_id:
            data[student_id][i] = {**d, **deed_data}
            found = True
            break
    if not found:
        data[student_id].append(deed_data)

    for json_p in [deeds_file, os.path.join(BASE_DIR, 'frontend', 'data', 'deeds.json')]:
        try:
            with open(json_p, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"⚠️ Error saving {json_p}: {e}")

    # 3. Update deeds_data.js in data/ and frontend/data/
    js_content = (
        f"// Auto-updated by server.py\n"
        f"const IMPORTED_DEEDS = {json.dumps(data, ensure_ascii=False, indent=2)};\n"
        f"const DEEDS_DATA = IMPORTED_DEEDS;\n\n"
        f"if (typeof window !== 'undefined') {{ window.IMPORTED_DEEDS = IMPORTED_DEEDS; window.DEEDS_DATA = DEEDS_DATA; }}\n"
        f"if (typeof globalThis !== 'undefined') {{ globalThis.IMPORTED_DEEDS = IMPORTED_DEEDS; globalThis.DEEDS_DATA = DEEDS_DATA; }}\n"
    )
    for js_p in [os.path.join(BASE_DIR, 'data', 'deeds_data.js'), os.path.join(BASE_DIR, 'frontend', 'data', 'deeds_data.js')]:
        try:
            with open(js_p, 'w', encoding='utf-8') as f:
                f.write(js_content)
        except Exception as e:
            print(f"⚠️ Error saving {js_p}: {e}")

    return True


class CustomHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=FRONTEND_DIR, **kwargs)

    def translate_path(self, path):
        if path.startswith('/frontend/'):
            path = path[9:]
        elif path == '/frontend':
            path = '/'
        return super().translate_path(path)

    def get_auth_context(self):
        cookies = parse_cookie_header(self.headers.get('Cookie'))
        return {
            'role': self.headers.get('X-GoodDeeds-Role') or cookies.get('gooddeeds_role') or '',
            'student_id': self.headers.get('X-GoodDeeds-Student-Id') or cookies.get('gooddeeds_student_id') or '',
            'username': self.headers.get('X-GoodDeeds-Username') or cookies.get('gooddeeds_username') or '',
        }

    def send_json_response(self, status_code, payload):
        self.send_response(status_code)
        self.send_header('Content-type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, Authorization, X-GoodDeeds-Role, X-GoodDeeds-Student-Id, X-GoodDeeds-Username')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.end_headers()
        self.wfile.write(json.dumps(payload, ensure_ascii=False).encode('utf-8'))

    def deny_json(self, message='Forbidden'):
        self.send_json_response(403, {'status': 'error', 'message': message})

    def can_access_student(self, student_id):
        auth = self.get_auth_context()
        if is_staff_role(auth.get('role')):
            return True
        return auth.get('role') == 'student' and auth.get('student_id') == str(student_id)

    def require_staff(self):
        if is_staff_role(self.get_auth_context().get('role')):
            return True
        self.deny_json('ต้องใช้สิทธิ์อาจารย์หรือผู้ดูแลระบบ')
        return False

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS, PUT, DELETE')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, Authorization, X-GoodDeeds-Role, X-GoodDeeds-Student-Id, X-GoodDeeds-Username')
        self.send_header('Access-Control-Max-Age', '86400')
        self.end_headers()

    def do_GET(self):
        parsed_path = urlparse(self.path)
        query_params = {}
        if parsed_path.query:
            query_params = {k: v[0] for k, v in parse_qs(parsed_path.query).items()}
            
        if parsed_path.path == '/api/get_student':
            student_id = query_params.get('studentId')
            if not student_id:
                self.send_json_response(400, {'status': 'error', 'message': 'Missing studentId'})
                return
            s_map = load_students_map()
            s = s_map.get(str(student_id).strip())
            if s:
                safe_s = {
                    'student_id': s.get('student_id'),
                    'rank': s.get('rank', 'นพอ.'),
                    'first_name': s.get('first_name', ''),
                    'last_name': s.get('last_name', ''),
                    'full_name': f"{s.get('rank', 'นพอ.')} {s.get('first_name', '')} {s.get('last_name', '')}".strip(),
                    'class_year': s.get('class_year', ''),
                    'year_level': s.get('year_level', ''),
                    'position': s.get('position', ''),
                    'role': s.get('role', 'student')
                }
                self.send_json_response(200, safe_s)
            else:
                self.send_json_response(404, {'status': 'error', 'message': 'Student not found'})
            return

        elif parsed_path.path == '/api/students':
            s_map = load_students_map()
            all_deeds = get_all_deeds()
            hours_map = {}
            for sid_k, d_list in all_deeds.items():
                tot = sum(float(d.get('hours', 0)) for d in d_list if d.get('status') == 'approved')
                hours_map[str(sid_k)] = round(tot, 1)

            safe_list = []
            for sid, s in s_map.items():
                sid_str = str(s.get('student_id'))
                safe_list.append({
                    'student_id': s.get('student_id'),
                    'rank': s.get('rank', 'นพอ.'),
                    'first_name': s.get('first_name', ''),
                    'last_name': s.get('last_name', ''),
                    'full_name': f"{s.get('rank', 'นพอ.')} {s.get('first_name', '')} {s.get('last_name', '')}".strip(),
                    'class_year': s.get('class_year', ''),
                    'year_level': s.get('year_level', ''),
                    'position': s.get('position', ''),
                    'role': s.get('role', 'student'),
                    'total_hours': hours_map.get(sid_str, 0)
                })
            self.send_json_response(200, safe_list)
            return

        elif parsed_path.path == '/api/get_deeds':
            student_id = query_params.get('studentId')
            if not student_id:
                self.send_response(400)
                self.end_headers()
                self.wfile.write(b"Missing studentId parameter")
                return
                
            try:
                deeds = get_deeds_for_student(student_id)
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(json.dumps(deeds, ensure_ascii=False).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.end_headers()
                self.wfile.write(str(e).encode('utf-8'))
                
        elif parsed_path.path == '/api/get_all_deeds':
            if not self.require_staff():
                return

            try:
                deeds = get_all_deeds()
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(json.dumps(deeds, ensure_ascii=False).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.end_headers()
                self.wfile.write(str(e).encode('utf-8'))
                
        elif parsed_path.path == '/api/export_docx':
            student_id = query_params.get('studentId')
            year_param = query_params.get('academicYear', '2569')
            try:
                academic_year = int(year_param)
            except ValueError:
                academic_year = 2569
                
            if not student_id:
                self.send_response(400)
                self.end_headers()
                self.wfile.write(b"Missing studentId parameter")
                return

            if not self.can_access_student(student_id):
                self.send_response(403)
                self.end_headers()
                self.wfile.write("นักเรียนสามารถส่งออกรายงานได้เฉพาะของตัวเองเท่านั้น".encode('utf-8'))
                return
                
            try:
                docx_buffer = generate_docx_in_memory(student_id, academic_year=academic_year)
                if not docx_buffer:
                    self.send_response(404)
                    self.end_headers()
                    self.wfile.write(b"Student not found")
                    return
                    
                self.send_response(200)
                self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Disposition', f'attachment; filename="report_{student_id}_{academic_year}.docx"')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(docx_buffer.getvalue())
            except Exception as e:
                self.send_response(500)
                self.end_headers()
                self.wfile.write(str(e).encode('utf-8'))
        elif parsed_path.path == '/api/events':
            self.send_response(200)
            self.send_header('Content-Type', 'text/event-stream')
            self.send_header('Cache-Control', 'no-cache')
            self.send_header('Connection', 'keep-alive')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            
            q = queue.Queue()
            with clients_lock:
                clients.append(q)
                
            try:
                while True:
                    try:
                        message = q.get(timeout=20.0)
                        self.wfile.write(message.encode('utf-8'))
                        self.wfile.flush()
                    except queue.Empty:
                        self.wfile.write(b": ping\n\n")
                        self.wfile.flush()
            except Exception:
                pass
            finally:
                with clients_lock:
                    if q in clients:
                        clients.remove(q)
            return
        else:
            # Serve static files normally (support both /frontend/path and /path)
            if self.path.startswith('/frontend/'):
                self.path = self.path[9:]
            super().do_GET()

    def do_POST(self):
        parsed_path = urlparse(self.path)
        
        if parsed_path.path == '/api/submit_deed':
            content_length = int(self.headers.get('Content-Length', 0))
            post_data = self.rfile.read(content_length)
            
            try:
                deed_data = json.loads(post_data.decode('utf-8'))
                student = deed_data.get('student', {})
                student_id = str(deed_data.get('studentId') or student.get('student_id') or '').strip()
                if not student_id:
                    self.send_json_response(400, {'status': 'error', 'message': 'Missing studentId'})
                    return

                deed_id = str(deed_data.get('id') or f"deed_{int(time.time()*1000)}_{student_id}")
                deed_data['id'] = deed_id
                deed_data['studentId'] = student_id
                deed_data['status'] = deed_data.get('status', 'pending')
                if 'submittedAt' not in deed_data:
                    deed_data['submittedAt'] = time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime())

                # Validate limits and prevent duplicates according to RTAFNC rules
                is_valid, validation_msg = validate_deed_submission(deed_data, student_id)
                if not is_valid:
                    self.send_json_response(400, {'status': 'error', 'message': validation_msg})
                    print(f"⚠️ Deed submission rejected for {student_id}: {validation_msg}")
                    return

                # Save evidence image to disk if base64 provided
                img_data = deed_data.get('imageData') or deed_data.get('imageUrl')
                if not img_data and isinstance(deed_data.get('imageUrls'), list) and len(deed_data['imageUrls']) > 0:
                    first_img = deed_data['imageUrls'][0]
                    if isinstance(first_img, str) and first_img.startswith('data:image'):
                        img_data = first_img

                if img_data and isinstance(img_data, str) and img_data.startswith('data:image'):
                    try:
                        import base64 as b64module
                        _, b64_content = img_data.split('base64,', 1)
                        img_bytes = b64module.b64decode(b64_content)
                        evidence_dir = os.path.join(FRONTEND_DIR, 'photos', 'evidence')
                        os.makedirs(evidence_dir, exist_ok=True)
                        img_filename = f"{deed_id}.jpg"
                        img_path = os.path.join(evidence_dir, img_filename)
                        with open(img_path, 'wb') as img_f:
                            img_f.write(img_bytes)
                        rel_path = f"photos/evidence/{img_filename}"
                        deed_data['imageUrl'] = rel_path
                        deed_data['imageUrls'] = [rel_path]
                        deed_data.pop('imageData', None)
                        print(f"📸 Saved deed evidence photo: {rel_path} ({len(img_bytes)} bytes)")
                    except Exception as img_err:
                        print(f"⚠️ Error saving evidence photo: {img_err}")

                save_or_update_deed_in_db(student_id, deed_data)

                self.send_json_response(200, {
                    'status': 'success',
                    'message': 'Deed submitted and saved successfully',
                    'deedId': deed_id
                })
                print(f"📥 Saved PENDING deed [{deed_id}] for student {student_id}")
                broadcast_event("deed_submitted", {"studentId": student_id, "deedId": deed_id})
                sync_to_google_drive_bg()
            except Exception as e:
                self.send_json_response(500, {'status': 'error', 'message': str(e)})
                print(f"❌ Error submitting deed: {e}")
                
        elif parsed_path.path == '/api/approve_deed':
            content_length = int(self.headers.get('Content-Length', 0))
            post_data = self.rfile.read(content_length)
            
            try:
                payload = json.loads(post_data.decode('utf-8'))
                student_id = str(payload.get('studentId') or '').strip()
                deed_id = str(payload.get('deedId') or '').strip()
                status = payload.get('status', 'approved')
                teacher_name = payload.get('teacherName', 'อาจารย์ผู้ตรวจประเมิน')
                reject_reason = payload.get('rejectReason', '')
                deed_data = payload.get('deedData') or {}

                if not student_id or not deed_id:
                    self.send_json_response(400, {'status': 'error', 'message': 'Missing studentId or deedId'})
                    return

                # If deed_data incomplete, fetch existing deed from DB
                existing_deeds = get_deeds_for_student(student_id)
                found = next((d for d in existing_deeds if str(d.get('id')) == deed_id), None)
                if found:
                    deed_data = {**found, **deed_data}

                deed_data['id'] = deed_id
                deed_data['studentId'] = student_id
                deed_data['status'] = status
                deed_data['approvedBy'] = teacher_name
                deed_data['approved_by'] = teacher_name
                deed_data['approvedAt'] = time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime())
                deed_data['updated_at'] = deed_data['approvedAt']
                if payload.get('signatureKey'):
                    deed_data['signatureKey'] = payload.get('signatureKey')
                if payload.get('signature'):
                    deed_data['signature'] = payload.get('signature')
                if status == 'rejected':
                    deed_data['rejectReason'] = reject_reason

                save_or_update_deed_in_db(student_id, deed_data)

                self.send_json_response(200, {
                    'status': 'success',
                    'message': f"Deed successfully updated to {status}",
                    'deedId': deed_id,
                    'status': status
                })
                print(f"❇️ Updated deed [{deed_id}] for {student_id} to {status} by {teacher_name}")
                broadcast_event("deed_approved", {"studentId": student_id, "deedId": deed_id, "status": status})
                sync_to_google_drive_bg()
            except Exception as e:
                self.send_json_response(500, {'status': 'error', 'message': str(e)})
                print(f"❌ Error updating deed: {e}")
                
        elif parsed_path.path == '/api/update_student':
            if not self.require_staff():
                return

            content_length = int(self.headers.get('Content-Length', 0))
            post_data = self.rfile.read(content_length)
            
            try:
                payload = json.loads(post_data.decode('utf-8'))
                student_id = payload.get('studentId')
                rank = payload.get('rank', 'นพอ.')
                first_name = payload.get('firstName', '')
                last_name = payload.get('lastName', '')
                class_year = int(payload.get('classYear', 69))
                email = payload.get('email', '')
                position = payload.get('position', 'นักเรียนพยาบาล')
                
                # Update in frontend/data/students.json
                frontend_json_path = os.path.join(BASE_DIR, 'frontend', 'data', 'students.json')
                frontend_js_path = os.path.join(BASE_DIR, 'frontend', 'data', 'students_data.js')
                root_json_path = os.path.join(BASE_DIR, 'data', 'students.json')
                root_js_path = os.path.join(BASE_DIR, 'data', 'students_data.js')
                
                students_list = []
                if os.path.exists(frontend_json_path):
                    with open(frontend_json_path, 'r', encoding='utf-8') as f:
                        students_list = json.load(f)
                
                # Find student and update
                updated = False
                for s in students_list:
                    if s.get('student_id') == student_id:
                        s['rank'] = rank
                        s['first_name'] = first_name
                        s['last_name'] = last_name
                        s['full_name'] = f"{first_name} {last_name}".strip()
                        s['class_year'] = class_year
                        s['email'] = email
                        s['position'] = position
                        if class_year == 69: s['year_level'] = 1
                        elif class_year == 68: s['year_level'] = 2
                        elif class_year == 67: s['year_level'] = 3
                        elif class_year == 66: s['year_level'] = 4
                        else: s['year_level'] = 5
                        updated = True
                        break
                        
                if not updated:
                    new_student = {
                        'student_id': student_id,
                        'rank': rank,
                        'first_name': first_name,
                        'last_name': last_name,
                        'full_name': f"{first_name} {last_name}".strip(),
                        'class_year': class_year,
                        'year_level': 1 if class_year == 69 else (2 if class_year == 68 else (3 if class_year == 67 else (4 if class_year == 66 else 5))),
                        'note': 'เพิ่มโดยอาจารย์ผ่านระบบ',
                        'password': student_id,
                        'email': email,
                        'telegram_chat_id': '',
                        'role': 'student'
                    }
                    students_list.append(new_student)
                
                # Sort
                students_list.sort(key=lambda x: (x.get('class_year', 69), x.get('student_id', '')))
                
                # Write back to frontend/data/students.json
                with open(frontend_json_path, 'w', encoding='utf-8') as f:
                    json.dump(students_list, f, ensure_ascii=False, indent=2)
                    
                # Write back to frontend/data/students_data.js
                with open(frontend_js_path, 'w', encoding='utf-8') as f:
                    f.write("// Auto-generated student data - DO NOT EDIT MANUALLY\n")
                    f.write("// Generated from: รายชื่อ นพอ.ปี69 ทุกชั้นปี\n\n")
                    f.write("const STUDENTS_DATA = ")
                    json.dump(students_list, f, ensure_ascii=False, indent=2)
                    f.write(";\n")
                    
                # Sync write to root data/ folder if paths exist
                if os.path.exists(os.path.dirname(root_json_path)):
                    with open(root_json_path, 'w', encoding='utf-8') as f:
                        json.dump(students_list, f, ensure_ascii=False, indent=2)
                    with open(root_js_path, 'w', encoding='utf-8') as f:
                        f.write("// Auto-generated student data - DO NOT EDIT MANUALLY\n")
                        f.write("const STUDENTS_DATA = ")
                        json.dump(students_list, f, ensure_ascii=False, indent=2)
                        f.write(";\n")
                
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                response = {
                    'status': 'success',
                    'message': 'Student updated successfully'
                }
                self.wfile.write(json.dumps(response).encode('utf-8'))
                print(f"👤 Updated student profile for {student_id}")
                broadcast_event("student_updated", {"studentId": student_id})
                
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                response = {'status': 'error', 'message': str(e)}
                self.wfile.write(json.dumps(response).encode('utf-8'))
                print(f"❌ Error updating student: {e}")
        elif self.path == '/api/bind_line':
            content_length = int(self.headers.get('Content-Length', 0))
            post_data = self.rfile.read(content_length)
            try:
                payload = json.loads(post_data.decode('utf-8'))
                student_id = payload.get('studentId')
                line_user_id = payload.get('lineUserId')
                line_name = payload.get('lineDisplayName', '')
                line_pic = payload.get('linePictureUrl', '')
                
                if student_id and line_user_id:
                    # 1. Store strictly in private backend storage (never in git or public frontend)
                    private_map_path = os.path.join(BASE_DIR, 'data', 'private', 'line_mappings.json')
                    mappings = {}
                    if os.path.exists(private_map_path):
                        try:
                            with open(private_map_path, 'r', encoding='utf-8') as f:
                                mappings = json.load(f)
                        except Exception:
                            mappings = {}
                    
                    mappings[str(student_id)] = {
                        'student_id': str(student_id),
                        'line_user_id': line_user_id,
                        'line_display_name': line_name,
                        'line_picture_url': line_pic,
                        'updated_at': time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime())
                    }
                    
                    with open(private_map_path, 'w', encoding='utf-8') as f:
                        json.dump(mappings, f, ensure_ascii=False, indent=2)

                    print(f"🔒 Bound LINE ID safely in private backend ({line_user_id[:8]}... - {line_name}) for student {student_id}")
                    sync_to_google_drive_bg()

                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'status': 'success'}).encode('utf-8'))
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'status': 'error', 'message': str(e)}).encode('utf-8'))
        else:
            self.send_response(404)
            self.end_headers()

def run(server_class=ThreadingHTTPServer, handler_class=CustomHandler, port=8000):
    server_address = ('', port)
    httpd = server_class(server_address, handler_class)
    print(f"🚀 Starting custom server on port {port}...")
    print(f"📂 Serving static files from {BASE_DIR}")
    print(f"📁 Saving records to {RECORDS_DIR}")
    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        pass
    httpd.server_close()
    print("Stopping server...\n")

if __name__ == '__main__':
    port = 8000
    if len(sys.argv) > 1:
        try:
            port = int(sys.argv[1])
        except ValueError:
            pass
    run(port=port)
